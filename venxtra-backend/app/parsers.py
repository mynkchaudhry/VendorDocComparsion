import pdfplumber
from docx import Document as DocxDocument
import openpyxl
from typing import Optional
import io

async def extract_text_from_pdf(file_content: bytes) -> str:
    text = []
    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    return "\n".join(text)

async def extract_text_from_docx(file_content: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_content))
    text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            if any(row_text):
                text.append(" | ".join(row_text))
    
    return "\n".join(text)

async def extract_text_from_excel(file_content: bytes) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(file_content), read_only=True)
    text = []
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        text.append(f"Sheet: {sheet_name}")
        
        for row in sheet.iter_rows(values_only=True):
            row_data = [str(cell) if cell is not None else "" for cell in row]
            if any(row_data):
                text.append(" | ".join(row_data))
    
    wb.close()
    return "\n".join(text)

async def extract_text(file_content: bytes, file_type: str) -> Optional[str]:
    try:
        if file_type.lower() in ['pdf', '.pdf']:
            return await extract_text_from_pdf(file_content)
        elif file_type.lower() in ['docx', '.docx', 'doc', '.doc']:
            return await extract_text_from_docx(file_content)
        elif file_type.lower() in ['xlsx', '.xlsx', 'xls', '.xls']:
            return await extract_text_from_excel(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    except Exception as e:
        raise Exception(f"Error extracting text: {str(e)}")