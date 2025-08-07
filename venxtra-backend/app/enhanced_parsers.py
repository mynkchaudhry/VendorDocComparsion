import pdfplumber
import PyPDF2
from docx import Document as DocxDocument
import openpyxl
import pandas as pd
import xlrd
from typing import Optional, List, Dict, Any, Generator, Union, Tuple
import io
import asyncio
import hashlib
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
import multiprocessing as mp
from dataclasses import dataclass, field
import math
import logging
import time
import sys
import gc
from pathlib import Path
import json
from enum import Enum
import warnings
warnings.filterwarnings("ignore", category=UserWarning)

logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """Represents a chunk of document content"""
    chunk_id: str
    page_range: str
    content: str
    word_count: int
    metadata: Dict[str, Any]
    tables: List[Dict[str, Any]] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)
    processing_time: Optional[float] = None
    quality_score: Optional[float] = None
    
    def __post_init__(self):
        """Calculate quality score based on content characteristics"""
        if self.content:
            # Simple quality metrics
            char_count = len(self.content)
            line_count = len(self.content.split('\n'))
            avg_line_length = char_count / line_count if line_count > 0 else 0
            
            # Quality score based on content characteristics
            self.quality_score = min(1.0, (char_count * 0.001 + avg_line_length * 0.01) / 2)

@dataclass
class ProcessingMetrics:
    """Comprehensive processing metrics"""
    start_time: float
    end_time: Optional[float] = None
    total_processing_time: Optional[float] = None
    memory_peak_mb: float = 0.0
    memory_average_mb: float = 0.0
    pages_processed: int = 0
    chunks_created: int = 0
    tables_extracted: int = 0
    images_extracted: int = 0
    errors_encountered: int = 0
    warnings_generated: int = 0
    avg_processing_speed_pages_per_sec: float = 0.0
    file_size_mb: float = 0.0

class ProcessingStatus(Enum):
    """Processing status enumeration"""
    INITIALIZING = "initializing"
    READING_FILE = "reading_file"
    EXTRACTING_TEXT = "extracting_text"
    CHUNKING = "chunking"
    PROCESSING_TABLES = "processing_tables"
    PROCESSING_IMAGES = "processing_images"
    FINALIZING = "finalizing"
    COMPLETED = "completed"
    ERROR = "error"

@dataclass
class ProcessingProgress:
    """Tracks document processing progress"""
    total_pages: int
    processed_pages: int
    total_chunks: int
    processed_chunks: int
    current_stage: str
    status: ProcessingStatus = ProcessingStatus.INITIALIZING
    estimated_time_remaining: Optional[float] = None
    memory_usage_mb: Optional[float] = None
    processing_speed_pages_per_sec: Optional[float] = None
    error_count: int = 0
    warnings: List[str] = field(default_factory=list)

class EnhancedDocumentProcessor:
    """Enhanced document processor with chunking, parallel processing, and advanced features"""
    
    def __init__(self, 
                 max_chunk_size: int = 2000,  # words per chunk
                 overlap_size: int = 200,     # word overlap between chunks
                 max_workers: int = None,
                 memory_limit_mb: int = 500,  # Memory limit for processing
                 enable_table_extraction: bool = True,
                 enable_image_extraction: bool = False,
                 quality_threshold: float = 0.1):
        self.max_chunk_size = max_chunk_size
        self.overlap_size = overlap_size
        self.max_workers = max_workers or min(4, mp.cpu_count())
        self.memory_limit_mb = memory_limit_mb
        self.enable_table_extraction = enable_table_extraction
        self.enable_image_extraction = enable_image_extraction
        self.quality_threshold = quality_threshold
        self.executor = ThreadPoolExecutor(max_workers=self.max_workers)
        self.metrics = ProcessingMetrics(start_time=time.time())
        
        # Performance tracking
        self._memory_samples = []
        self._processing_times = []
    
    def _get_memory_usage(self) -> float:
        """Get current memory usage in MB"""
        import psutil
        process = psutil.Process()
        return process.memory_info().rss / 1024 / 1024
    
    def _monitor_memory(self):
        """Monitor memory usage during processing"""
        current_memory = self._get_memory_usage()
        self._memory_samples.append(current_memory)
        
        if current_memory > self.metrics.memory_peak_mb:
            self.metrics.memory_peak_mb = current_memory
        
        # Check memory limit
        if current_memory > self.memory_limit_mb:
            gc.collect()  # Force garbage collection
            logger.warning(f"Memory usage ({current_memory:.1f}MB) exceeded limit ({self.memory_limit_mb}MB)")
    
    def _calculate_chunk_hash(self, content: str) -> str:
        """Generate a unique hash for chunk content"""
        return hashlib.md5(content.encode()).hexdigest()[:12]
    
    def _extract_tables_with_pandas(self, file_content: bytes) -> List[Dict[str, Any]]:
        """Extract tables using pandas for better performance"""
        tables = []
        try:
            # Try to read Excel with pandas
            dfs = pd.read_excel(io.BytesIO(file_content), sheet_name=None, engine='openpyxl')
            for sheet_name, df in dfs.items():
                if not df.empty:
                    tables.append({
                        'sheet_name': sheet_name,
                        'data': df.to_dict('records'),
                        'shape': df.shape,
                        'columns': df.columns.tolist()
                    })
        except Exception as e:
            logger.warning(f"Failed to extract tables with pandas: {str(e)}")
        return tables
    
    def _split_text_into_chunks(self, text: str, page_info: str, 
                               tables: List[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """Split text into overlapping chunks with enhanced metadata"""
        start_time = time.time()
        words = text.split()
        chunks = []
        tables = tables or []
        
        # Filter out low-quality content
        if len(text.strip()) < 10:  # Skip very short content
            return chunks
        
        if len(words) <= self.max_chunk_size:
            # Single chunk
            chunk_id = self._calculate_chunk_hash(text)
            chunk = DocumentChunk(
                chunk_id=chunk_id,
                page_range=page_info,
                content=text,
                word_count=len(words),
                metadata={
                    'is_single_chunk': True,
                    'extraction_method': 'single_chunk',
                    'language_detected': self._detect_language(text[:100])
                },
                tables=tables,
                processing_time=time.time() - start_time
            )
            # Only add if quality is acceptable
            if chunk.quality_score and chunk.quality_score >= self.quality_threshold:
                chunks.append(chunk)
        else:
            # Multiple chunks with overlap
            start = 0
            chunk_num = 0
            
            while start < len(words):
                end = min(start + self.max_chunk_size, len(words))
                chunk_words = words[start:end]
                chunk_text = ' '.join(chunk_words)
                
                chunk_id = f"{self._calculate_chunk_hash(chunk_text)}_{chunk_num}"
                chunk = DocumentChunk(
                    chunk_id=chunk_id,
                    page_range=f"{page_info}_chunk_{chunk_num}",
                    content=chunk_text,
                    word_count=len(chunk_words),
                    metadata={
                        'chunk_number': chunk_num,
                        'start_word': start,
                        'end_word': end,
                        'total_chunks_in_page': None,  # Will be updated later
                        'extraction_method': 'multi_chunk',
                        'language_detected': self._detect_language(chunk_text[:100])
                    },
                    tables=tables if chunk_num == 0 else [],  # Tables only in first chunk
                    processing_time=time.time() - start_time
                )
                
                # Only add if quality is acceptable
                if chunk.quality_score and chunk.quality_score >= self.quality_threshold:
                    chunks.append(chunk)
                
                # Move start position with overlap
                start = end - self.overlap_size if end < len(words) else end
                chunk_num += 1
            
            # Update total chunks info
            for chunk in chunks:
                chunk.metadata['total_chunks_in_page'] = len(chunks)
        
        return chunks
    
    def _detect_language(self, text_sample: str) -> str:
        """Simple language detection based on common patterns"""
        # Very basic language detection - can be enhanced with proper libraries
        if not text_sample:
            return 'unknown'
        
        # Count common English words
        english_indicators = ['the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with']
        words = text_sample.lower().split()
        english_count = sum(1 for word in words if word in english_indicators)
        
        if len(words) > 0 and english_count / len(words) > 0.1:
            return 'english'
        return 'other'
    
    async def extract_text_from_pdf_chunked(self, file_content: bytes, 
                                          progress_callback=None) -> Tuple[List[DocumentChunk], ProcessingMetrics]:
        """Extract text from PDF with enhanced chunking and progress tracking"""
        all_chunks = []
        self.metrics = ProcessingMetrics(start_time=time.time())
        self.metrics.file_size_mb = len(file_content) / 1024 / 1024
        
        try:
            self._monitor_memory()
            
            # Try multiple PDF processing approaches for robustness
            pdf_reader = None
            total_pages = 0
            
            # First try pdfplumber (better for text and tables)
            try:
                pdf_reader = pdfplumber.open(io.BytesIO(file_content))
                total_pages = len(pdf_reader.pages)
                logger.info(f"Using pdfplumber for PDF processing ({total_pages} pages)")
            except Exception as e:
                logger.warning(f"pdfplumber failed: {str(e)}, trying PyPDF2")
                
                # Fallback to PyPDF2
                try:
                    pdf_bytes = io.BytesIO(file_content)
                    pdf_reader = PyPDF2.PdfReader(pdf_bytes)
                    total_pages = len(pdf_reader.pages)
                    logger.info(f"Using PyPDF2 for PDF processing ({total_pages} pages)")
                except Exception as e2:
                    raise Exception(f"Both PDF readers failed: pdfplumber: {str(e)}, PyPDF2: {str(e2)}")
            
            if progress_callback:
                await progress_callback(ProcessingProgress(
                    total_pages=total_pages,
                    processed_pages=0,
                    total_chunks=0,
                    processed_chunks=0,
                    current_stage="Extracting text from pages",
                    status=ProcessingStatus.EXTRACTING_TEXT,
                    memory_usage_mb=self._get_memory_usage()
                ))
                
            # Process pages in batches for better memory management
            batch_size = max(1, min(10, total_pages))  # Process 1-10 pages at a time
            batch_start_time = time.time()
            
            for batch_start in range(0, total_pages, batch_size):
                batch_end = min(batch_start + batch_size, total_pages)
                batch_chunks = []
                
                # Memory management
                self._monitor_memory()
                
                # Process batch based on PDF reader type
                if isinstance(pdf_reader, pdfplumber.PDF):
                    batch_pages = pdf_reader.pages[batch_start:batch_end]
                    
                    for i, page in enumerate(batch_pages):
                        page_num = batch_start + i + 1
                        page_start_time = time.time()
                        
                        try:
                            page_text = page.extract_text() or ""
                            tables_data = []
                            
                            if page_text.strip():
                                # Extract tables if enabled
                                if self.enable_table_extraction:
                                    try:
                                        tables = page.extract_tables()
                                        if tables:
                                            for table_idx, table in enumerate(tables):
                                                if table:
                                                    table_data = []
                                                    for row in table:
                                                        if row:
                                                            clean_row = [str(cell).strip() if cell else "" for cell in row]
                                                            if any(clean_row):
                                                                table_data.append(clean_row)
                                                    
                                                    if table_data:
                                                        tables_data.append({
                                                            'table_id': f"page_{page_num}_table_{table_idx}",
                                                            'data': table_data,
                                                            'page': page_num
                                                        })
                                                        # Add table text to content
                                                        table_text = "\n".join([" | ".join(row) for row in table_data])
                                                        page_text += f"\n\nTable {table_idx + 1}:\n{table_text}"
                                                        self.metrics.tables_extracted += 1
                                    except Exception as e:
                                        logger.warning(f"Table extraction failed for page {page_num}: {str(e)}")
                                        self.metrics.warnings_generated += 1
                                
                                # Create chunks for this page
                                page_chunks = self._split_text_into_chunks(
                                    page_text, 
                                    f"page_{page_num}",
                                    tables_data
                                )
                                batch_chunks.extend(page_chunks)
                                
                                # Track processing time
                                page_time = time.time() - page_start_time
                                self._processing_times.append(page_time)
                                
                        except Exception as e:
                            logger.error(f"Error processing page {page_num}: {str(e)}")
                            self.metrics.errors_encountered += 1
                else:
                    # Handle PyPDF2
                    for i in range(batch_start, batch_end):
                        page_num = i + 1
                        try:
                            page = pdf_reader.pages[i]
                            page_text = page.extract_text() or ""
                            
                            if page_text.strip():
                                page_chunks = self._split_text_into_chunks(
                                    page_text, 
                                    f"page_{page_num}"
                                )
                                batch_chunks.extend(page_chunks)
                        except Exception as e:
                            logger.error(f"Error processing page {page_num} with PyPDF2: {str(e)}")
                            self.metrics.errors_encountered += 1
                        
                # Update progress
                processed_pages = min(batch_end, total_pages)
                if progress_callback:
                    # Calculate processing speed
                    elapsed_time = time.time() - batch_start_time
                    speed = processed_pages / elapsed_time if elapsed_time > 0 else 0
                    
                    await progress_callback(ProcessingProgress(
                        total_pages=total_pages,
                        processed_pages=processed_pages,
                        total_chunks=len(all_chunks) + len(batch_chunks),
                        processed_chunks=len(all_chunks),
                        current_stage=f"Processing page {processed_pages}/{total_pages}",
                        status=ProcessingStatus.EXTRACTING_TEXT,
                        memory_usage_mb=self._get_memory_usage(),
                        processing_speed_pages_per_sec=speed,
                        error_count=self.metrics.errors_encountered
                    ))
                
                all_chunks.extend(batch_chunks)
                self.metrics.pages_processed = processed_pages
                self.metrics.chunks_created = len(all_chunks)
                    
                # Small delay and cleanup to prevent memory issues
                if batch_end < total_pages:
                    await asyncio.sleep(0.1)
                    gc.collect()  # Force garbage collection between batches
                
            # Finalize processing
            self.metrics.end_time = time.time()
            self.metrics.total_processing_time = self.metrics.end_time - self.metrics.start_time
            
            # Calculate averages
            if self._memory_samples:
                self.metrics.memory_average_mb = sum(self._memory_samples) / len(self._memory_samples)
            
            if self.metrics.total_processing_time > 0:
                self.metrics.avg_processing_speed_pages_per_sec = total_pages / self.metrics.total_processing_time
            
            if progress_callback:
                await progress_callback(ProcessingProgress(
                    total_pages=total_pages,
                    processed_pages=total_pages,
                    total_chunks=len(all_chunks),
                    processed_chunks=len(all_chunks),
                    current_stage="Text extraction completed",
                    status=ProcessingStatus.COMPLETED,
                    memory_usage_mb=self._get_memory_usage(),
                    processing_speed_pages_per_sec=self.metrics.avg_processing_speed_pages_per_sec,
                    error_count=self.metrics.errors_encountered
                ))
                
            logger.info(f"Successfully extracted {len(all_chunks)} chunks from {total_pages} pages")
            logger.info(f"Processing metrics: {self.metrics.total_processing_time:.2f}s, "
                       f"{self.metrics.avg_processing_speed_pages_per_sec:.2f} pages/sec, "
                       f"Peak memory: {self.metrics.memory_peak_mb:.1f}MB")
            
            return all_chunks, self.metrics
                
        except Exception as e:
            self.metrics.end_time = time.time()
            self.metrics.errors_encountered += 1
            logger.error(f"Error processing PDF: {str(e)}")
            raise Exception(f"PDF processing failed: {str(e)}")
        finally:
            # Cleanup
            if pdf_reader and hasattr(pdf_reader, 'close'):
                try:
                    pdf_reader.close()
                except:
                    pass
            gc.collect()
    
    async def extract_text_from_docx_chunked(self, file_content: bytes,
                                           progress_callback=None) -> Tuple[List[DocumentChunk], ProcessingMetrics]:
        """Extract text from DOCX with chunking"""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    if any(row_text):
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n".join(text_parts)
            
            if progress_callback:
                await progress_callback(ProcessingProgress(
                    total_pages=1,
                    processed_pages=1,
                    total_chunks=0,
                    processed_chunks=0,
                    current_stage="Processing DOCX content"
                ))
            
            chunks = self._split_text_into_chunks(full_text, "docx_content")
            
            if progress_callback:
                await progress_callback(ProcessingProgress(
                    total_pages=1,
                    processed_pages=1,
                    total_chunks=len(chunks),
                    processed_chunks=len(chunks),
                    current_stage="DOCX processing completed"
                ))
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise Exception(f"DOCX processing failed: {str(e)}")
    
    async def extract_text_from_excel_chunked(self, file_content: bytes,
                                            progress_callback=None) -> Tuple[List[DocumentChunk], ProcessingMetrics]:
        """Extract text from Excel with enhanced chunking and pandas support"""
        self.metrics = ProcessingMetrics(start_time=time.time())
        self.metrics.file_size_mb = len(file_content) / 1024 / 1024
        all_chunks = []
        
        try:
            self._monitor_memory()
            
            # Try pandas first for better performance with large files
            excel_data = None
            processing_method = "pandas"
            
            try:
                # Use pandas for better memory management
                excel_data = pd.read_excel(
                    io.BytesIO(file_content), 
                    sheet_name=None,
                    engine='openpyxl' if file_content[:4] == b'PK\x03\x04' else 'xlrd'
                )
                logger.info(f"Using pandas for Excel processing ({len(excel_data)} sheets)")
            except Exception as e:
                logger.warning(f"Pandas failed: {str(e)}, falling back to openpyxl")
                processing_method = "openpyxl"
            
            if excel_data is not None:
                # Process with pandas
                total_sheets = len(excel_data)
                processed_sheets = 0
                
                if progress_callback:
                    await progress_callback(ProcessingProgress(
                        total_pages=total_sheets,
                        processed_pages=0,
                        total_chunks=0,
                        processed_chunks=0,
                        current_stage="Processing Excel sheets with pandas",
                        status=ProcessingStatus.EXTRACTING_TEXT,
                        memory_usage_mb=self._get_memory_usage()
                    ))
                
                for sheet_name, df in excel_data.items():
                    sheet_start_time = time.time()
                    self._monitor_memory()
                    
                    try:
                        if not df.empty:
                            # Convert DataFrame to structured text
                            sheet_parts = [f"Sheet: {sheet_name}"]
                            
                            # Add column headers
                            headers = df.columns.tolist()
                            sheet_parts.append("Headers: " + " | ".join([str(h) for h in headers]))
                            
                            # Process rows in batches to manage memory
                            batch_size = min(1000, len(df))
                            for batch_start in range(0, len(df), batch_size):
                                batch_end = min(batch_start + batch_size, len(df))
                                batch_df = df.iloc[batch_start:batch_end]
                                
                                for _, row in batch_df.iterrows():
                                    row_data = [str(cell) if pd.notna(cell) else "" for cell in row]
                                    if any(row_data):
                                        sheet_parts.append(" | ".join(row_data))
                            
                            sheet_content = "\n".join(sheet_parts)
                            
                            # Create table metadata
                            tables_data = [{
                                'sheet_name': sheet_name,
                                'shape': df.shape,
                                'columns': headers,
                                'data_types': df.dtypes.to_dict(),
                                'non_null_counts': df.count().to_dict()
                            }]
                            
                            # Create chunks for this sheet
                            sheet_chunks = self._split_text_into_chunks(
                                sheet_content,
                                f"sheet_{sheet_name}",
                                tables_data
                            )
                            all_chunks.extend(sheet_chunks)
                            self.metrics.tables_extracted += 1
                            
                            # Track processing time
                            sheet_time = time.time() - sheet_start_time
                            self._processing_times.append(sheet_time)
                            
                    except Exception as e:
                        logger.error(f"Error processing sheet {sheet_name}: {str(e)}")
                        self.metrics.errors_encountered += 1
                    
                    processed_sheets += 1
                    
                    if progress_callback:
                        await progress_callback(ProcessingProgress(
                            total_pages=total_sheets,
                            processed_pages=processed_sheets,
                            total_chunks=len(all_chunks),
                            processed_chunks=len(all_chunks),
                            current_stage=f"Processing sheet {processed_sheets}/{total_sheets}",
                            status=ProcessingStatus.EXTRACTING_TEXT,
                            memory_usage_mb=self._get_memory_usage(),
                            error_count=self.metrics.errors_encountered
                        ))
            
            else:
                # Fallback to openpyxl
                wb = openpyxl.load_workbook(io.BytesIO(file_content), read_only=True, data_only=True)
                total_sheets = len(wb.sheetnames)
                processed_sheets = 0
                
                if progress_callback:
                    await progress_callback(ProcessingProgress(
                        total_pages=total_sheets,
                        processed_pages=0,
                        total_chunks=0,
                        processed_chunks=0,
                        current_stage="Processing Excel sheets with openpyxl",
                        status=ProcessingStatus.EXTRACTING_TEXT,
                        memory_usage_mb=self._get_memory_usage()
                    ))
                
                for sheet_name in wb.sheetnames:
                    sheet_start_time = time.time()
                    self._monitor_memory()
                    
                    try:
                        sheet = wb[sheet_name]
                        sheet_text = [f"Sheet: {sheet_name}"]
                        
                        # Process rows in batches
                        row_batch = []
                        batch_size = 100
                        
                        for row in sheet.iter_rows(values_only=True):
                            row_data = [str(cell) if cell is not None else "" for cell in row]
                            if any(row_data):
                                row_batch.append(" | ".join(row_data))
                            
                            if len(row_batch) >= batch_size:
                                sheet_text.extend(row_batch)
                                row_batch = []
                        
                        # Add remaining rows
                        if row_batch:
                            sheet_text.extend(row_batch)
                        
                        sheet_content = "\n".join(sheet_text)
                        sheet_chunks = self._split_text_into_chunks(sheet_content, f"sheet_{sheet_name}")
                        all_chunks.extend(sheet_chunks)
                        
                        # Track processing time
                        sheet_time = time.time() - sheet_start_time
                        self._processing_times.append(sheet_time)
                        
                    except Exception as e:
                        logger.error(f"Error processing sheet {sheet_name}: {str(e)}")
                        self.metrics.errors_encountered += 1
                    
                    processed_sheets += 1
                    
                    if progress_callback:
                        await progress_callback(ProcessingProgress(
                            total_pages=total_sheets,
                            processed_pages=processed_sheets,
                            total_chunks=len(all_chunks),
                            processed_chunks=len(all_chunks),
                            current_stage=f"Processing sheet {processed_sheets}/{total_sheets}",
                            status=ProcessingStatus.EXTRACTING_TEXT,
                            memory_usage_mb=self._get_memory_usage(),
                            error_count=self.metrics.errors_encountered
                        ))
                
                wb.close()
            
            # Finalize processing
            self.metrics.end_time = time.time()
            self.metrics.total_processing_time = self.metrics.end_time - self.metrics.start_time
            self.metrics.chunks_created = len(all_chunks)
            
            # Calculate averages
            if self._memory_samples:
                self.metrics.memory_average_mb = sum(self._memory_samples) / len(self._memory_samples)
            
            if progress_callback:
                await progress_callback(ProcessingProgress(
                    total_pages=total_sheets if 'total_sheets' in locals() else 0,
                    processed_pages=total_sheets if 'total_sheets' in locals() else 0,
                    total_chunks=len(all_chunks),
                    processed_chunks=len(all_chunks),
                    current_stage=f"Excel processing completed using {processing_method}",
                    status=ProcessingStatus.COMPLETED,
                    memory_usage_mb=self._get_memory_usage(),
                    error_count=self.metrics.errors_encountered
                ))
            
            logger.info(f"Successfully extracted {len(all_chunks)} chunks from Excel file using {processing_method}")
            return all_chunks, self.metrics
            
        except Exception as e:
            self.metrics.end_time = time.time()
            self.metrics.errors_encountered += 1
            logger.error(f"Error processing Excel: {str(e)}")
            raise Exception(f"Excel processing failed: {str(e)}")
        finally:
            gc.collect()
        """Extract text from Excel with chunking"""
        try:
            wb = openpyxl.load_workbook(io.BytesIO(file_content), read_only=True)
            all_chunks = []
            
            total_sheets = len(wb.sheetnames)
            processed_sheets = 0
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_text = [f"Sheet: {sheet_name}"]
                
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_data):
                        sheet_text.append(" | ".join(row_data))
                
                sheet_content = "\n".join(sheet_text)
                sheet_chunks = self._split_text_into_chunks(sheet_content, f"sheet_{sheet_name}")
                all_chunks.extend(sheet_chunks)
                
                processed_sheets += 1
                
                if progress_callback:
                    await progress_callback(ProcessingProgress(
                        total_pages=total_sheets,
                        processed_pages=processed_sheets,
                        total_chunks=len(all_chunks),
                        processed_chunks=len(all_chunks),
                        current_stage=f"Processing sheet {processed_sheets}/{total_sheets}"
                    ))
            
            wb.close()
            return all_chunks
            
        except Exception as e:
            logger.error(f"Error processing Excel: {str(e)}")
            raise Exception(f"Excel processing failed: {str(e)}")

# Enhanced text extraction function with chunking
async def extract_text_chunked(file_content: bytes, file_type: str, 
                             progress_callback=None) -> List[DocumentChunk]:
    """Extract text from various file types with chunking support"""
    processor = EnhancedPDFProcessor()
    
    try:
        if file_type.lower() in ['pdf', '.pdf']:
            return await processor.extract_text_from_pdf_chunked(file_content, progress_callback)
        elif file_type.lower() in ['docx', '.docx', 'doc', '.doc']:
            return await processor.extract_text_from_docx_chunked(file_content, progress_callback)
        elif file_type.lower() in ['xlsx', '.xlsx', 'xls', '.xls']:
            return await processor.extract_text_from_excel_chunked(file_content, progress_callback)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    except Exception as e:
        logger.error(f"Error in enhanced text extraction: {str(e)}")
        raise Exception(f"Enhanced text extraction failed: {str(e)}")

# Backward compatibility function
async def extract_text_enhanced(file_content: bytes, file_type: str) -> str:
    """Enhanced text extraction that combines chunks into single text (backward compatible)"""
    chunks = await extract_text_chunked(file_content, file_type)
    return "\n\n".join([chunk.content for chunk in chunks])